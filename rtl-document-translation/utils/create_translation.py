#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Translation Dictionary Creator

Extracts unique text snippets from DOCX files and creates
translation dictionary template for RTL document translation.

Usage:
    python create_translation.py english1.docx english2.docx --output translations.json
"""

import json
import sys
import re
import argparse
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')


def normalize_text(text):
    """Normalize quotes and unicode spaces"""
    # Convert curly quotes to straight quotes
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2018', "'").replace('\u2019', "'")

    # Normalize unicode spaces
    text = re.sub(r'[\u2002\u2003\u2009\u200A\u00A0]+', ' ', text)

    return text.strip()


def extract_unique_texts(docx_paths):
    """Extract all unique text snippets from documents"""
    unique_texts = set()

    for docx_path in docx_paths:
        print(f"Processing: {docx_path}")
        doc = Document(docx_path)

        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                unique_texts.add(text)

                # Also add normalized version
                normalized = normalize_text(text)
                if normalized != text:
                    unique_texts.add(normalized)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        unique_texts.add(text)

                        normalized = normalize_text(text)
                        if normalized != text:
                            unique_texts.add(normalized)

    print(f"\nExtracted {len(unique_texts)} unique text snippets")
    return unique_texts


def create_translation_dictionary(unique_texts, output_path, target_lang='ar'):
    """
    Create translation dictionary

    For now, creates template with [TRANSLATE] placeholders.
    Can be extended to call translation APIs.
    """
    translations = {}

    # Self-map single letters (allowed English)
    single_letters = set('ABCDEFGHIS')

    # Self-map numbers and scoring notation
    number_pattern = re.compile(r'^-?\d+$|^\(-?\d+\)$')

    # Known abbreviations (self-map)
    abbreviations = {'DVD', 'CD', 'TV', 'USB', 'PDF', 'CEO', 'CFO'}

    for text in sorted(unique_texts):
        # Single letters
        if text in single_letters:
            translations[text] = text
            continue

        # Numbers and scoring
        if number_pattern.match(text):
            translations[text] = text
            continue

        # Abbreviations
        if text.upper() in abbreviations:
            translations[text] = text
            continue

        # Everything else needs translation
        # Option 1: Placeholder
        translations[text] = f"[TRANSLATE: {text}]"

        # Option 2: Call translation API (uncomment to use)
        # try:
        #     from googletrans import Translator
        #     translator = Translator()
        #     result = translator.translate(text, dest=target_lang)
        #     translations[text] = result.text
        # except Exception as e:
        #     print(f"Translation failed for '{text}': {e}")
        #     translations[text] = f"[TRANSLATE: {text}]"

    # Save
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(translations, f, ensure_ascii=False, indent=2, sort_keys=True)

    print(f"\nSaved {len(translations)} entries to {output_path}")
    print(f"\nNext steps:")
    print(f"1. Open {output_path}")
    print(f"2. Replace [TRANSLATE: ...] placeholders with actual translations")
    print(f"3. Or use translation API (uncomment code in script)")

    return translations


def main():
    parser = argparse.ArgumentParser(
        description='Create translation dictionary from DOCX files'
    )
    parser.add_argument(
        'docx_files',
        nargs='+',
        help='One or more DOCX files to process'
    )
    parser.add_argument(
        '--output', '-o',
        default='translation_dictionary.json',
        help='Output JSON file path (default: translation_dictionary.json)'
    )
    parser.add_argument(
        '--target-lang', '-t',
        default='ar',
        help='Target language code (default: ar for Arabic)'
    )

    args = parser.parse_args()

    print("="*80)
    print("TRANSLATION DICTIONARY CREATOR")
    print("="*80)

    # Extract unique texts
    unique_texts = extract_unique_texts(args.docx_files)

    # Create translation dictionary
    create_translation_dictionary(
        unique_texts,
        args.output,
        args.target_lang
    )

    print("\n" + "="*80)
    print("COMPLETE")
    print("="*80)


if __name__ == '__main__':
    main()
