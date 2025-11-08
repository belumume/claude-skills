#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
RTL Document Verification Utility

Comprehensive verification of RTL-translated DOCX files.

Usage:
    python verify_document.py arabic.docx english.docx translations.json
"""

import json
import sys
import re
import argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')


def load_allowed_english(translation_dict):
    """Extract allowed English from self-mapped dictionary entries"""
    allowed = set()

    for key, value in translation_dict.items():
        if key == value and re.search(r'[a-zA-Z]', key):
            allowed.add(key.strip())

    # Known abbreviations
    allowed.update(['DVD', 'CD', 'TV', 'USB', 'PDF', 'CEO', 'CFO'])

    return allowed


def scan_for_unauthorized_english(doc, allowed_english):
    """
    Scan document for unauthorized English words
    Returns list of (location, word) tuples
    """
    unauthorized = []

    def check_text(text, location):
        """Check text for English words not in allowed set"""
        # Remove numbers and scoring notation
        cleaned = re.sub(r'\(-?\d+\)', '', text)
        cleaned = re.sub(r'[$€£¥]?\d+[,.]?\d*%?', '', cleaned)
        cleaned = re.sub(r'\d+', '', cleaned)

        # Extract words
        words = re.findall(r'\b[a-zA-Z]+\b', cleaned)

        for word in words:
            if word not in allowed_english and word.upper() not in allowed_english:
                unauthorized.append((location, word))

    # Scan paragraphs
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            check_text(para.text, f"Paragraph {idx + 1}")

    # Scan tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    check_text(
                        cell.text,
                        f"Table {t_idx + 1}, Row {r_idx + 1}, Col {c_idx + 1}"
                    )

    return unauthorized


def verify_structure(ar_doc, eng_doc):
    """Verify structure matches between Arabic and English"""
    results = []

    # Sections
    if len(ar_doc.sections) != len(eng_doc.sections):
        results.append({
            'check': 'Sections',
            'status': 'FAIL',
            'expected': len(eng_doc.sections),
            'actual': len(ar_doc.sections)
        })
    else:
        results.append({
            'check': 'Sections',
            'status': 'PASS',
            'value': len(ar_doc.sections)
        })

    # Tables
    if len(ar_doc.tables) != len(eng_doc.tables):
        results.append({
            'check': 'Tables',
            'status': 'FAIL',
            'expected': len(eng_doc.tables),
            'actual': len(ar_doc.tables)
        })
    else:
        results.append({
            'check': 'Tables',
            'status': 'PASS',
            'value': len(ar_doc.tables)
        })

        # Check each table's dimensions
        for idx, (ar_table, eng_table) in enumerate(zip(ar_doc.tables, eng_doc.tables)):
            ar_rows = len(ar_table.rows)
            ar_cols = len(ar_table.columns)
            eng_rows = len(eng_table.rows)
            eng_cols = len(eng_table.columns)

            if ar_rows != eng_rows or ar_cols != eng_cols:
                results.append({
                    'check': f'Table {idx + 1} dimensions',
                    'status': 'FAIL',
                    'expected': f'{eng_rows}×{eng_cols}',
                    'actual': f'{ar_rows}×{ar_cols}'
                })
            else:
                results.append({
                    'check': f'Table {idx + 1} dimensions',
                    'status': 'PASS',
                    'value': f'{ar_rows}×{ar_cols}'
                })

    # Paragraphs
    if len(ar_doc.paragraphs) != len(eng_doc.paragraphs):
        results.append({
            'check': 'Paragraphs',
            'status': 'FAIL',
            'expected': len(eng_doc.paragraphs),
            'actual': len(ar_doc.paragraphs)
        })
    else:
        results.append({
            'check': 'Paragraphs',
            'status': 'PASS',
            'value': len(ar_doc.paragraphs)
        })

    return results


def verify_alignment(doc):
    """Verify all cells are right-aligned"""
    total_cells = 0
    right_aligned = 0
    misaligned = []

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                total_cells += 1

                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        right_aligned += 1
                    else:
                        misaligned.append(f"Table {t_idx + 1}, Row {r_idx + 1}, Col {c_idx + 1}")

    percentage = (right_aligned / total_cells * 100) if total_cells > 0 else 0

    return {
        'total': total_cells,
        'right_aligned': right_aligned,
        'percentage': percentage,
        'misaligned': misaligned[:10]  # First 10 only
    }


def verify_rtl_formatting(doc):
    """Verify RTL (bidi) formatting applied"""
    total_paragraphs = 0
    rtl_formatted = 0
    missing_rtl = []

    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            total_paragraphs += 1

            if para.paragraph_format.bidi:
                rtl_formatted += 1
            else:
                missing_rtl.append(f"Paragraph {idx + 1}")

    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.paragraphs and cell.text.strip():
                    total_paragraphs += 1

                    if cell.paragraphs[0].paragraph_format.bidi:
                        rtl_formatted += 1
                    else:
                        missing_rtl.append(f"Table {t_idx + 1}, Row {r_idx + 1}, Col {c_idx + 1}")

    percentage = (rtl_formatted / total_paragraphs * 100) if total_paragraphs > 0 else 0

    return {
        'total': total_paragraphs,
        'rtl_formatted': rtl_formatted,
        'percentage': percentage,
        'missing_rtl': missing_rtl[:10]
    }


def main():
    parser = argparse.ArgumentParser(
        description='Verify RTL-translated DOCX document'
    )
    parser.add_argument('arabic_docx', help='Arabic DOCX file to verify')
    parser.add_argument('english_docx', help='English DOCX file (reference)')
    parser.add_argument('translation_dict', help='Translation dictionary JSON')

    args = parser.parse_args()

    print("="*80)
    print("RTL DOCUMENT VERIFICATION")
    print("="*80)

    # Load documents
    print(f"\n[1/6] Loading documents...")
    ar_doc = Document(args.arabic_docx)
    eng_doc = Document(args.english_docx)

    with open(args.translation_dict, 'r', encoding='utf-8') as f:
        translation_dict = json.load(f)

    print(f"   Arabic: {args.arabic_docx}")
    print(f"   English: {args.english_docx}")
    print(f"   Dictionary: {len(translation_dict)} entries")

    # Check 1: Structure
    print(f"\n[2/6] Verifying structure...")
    structure_results = verify_structure(ar_doc, eng_doc)

    for result in structure_results:
        status_icon = '✓' if result['status'] == 'PASS' else '✗'
        if result['status'] == 'PASS':
            print(f"   {status_icon} {result['check']}: {result['value']}")
        else:
            print(f"   {status_icon} {result['check']}: Expected {result['expected']}, Got {result['actual']}")

    # Check 2: Alignment
    print(f"\n[3/6] Verifying alignment...")
    alignment_result = verify_alignment(ar_doc)

    print(f"   Total cells: {alignment_result['total']}")
    print(f"   Right-aligned: {alignment_result['right_aligned']}/{alignment_result['total']} ({alignment_result['percentage']:.1f}%)")

    if alignment_result['misaligned']:
        print(f"   ✗ Misaligned cells (first 10):")
        for loc in alignment_result['misaligned']:
            print(f"      - {loc}")

    # Check 3: RTL Formatting
    print(f"\n[4/6] Verifying RTL formatting...")
    rtl_result = verify_rtl_formatting(ar_doc)

    print(f"   Total elements: {rtl_result['total']}")
    print(f"   RTL formatted: {rtl_result['rtl_formatted']}/{rtl_result['total']} ({rtl_result['percentage']:.1f}%)")

    if rtl_result['missing_rtl']:
        print(f"   ✗ Missing RTL (first 10):")
        for loc in rtl_result['missing_rtl']:
            print(f"      - {loc}")

    # Check 4: English Words
    print(f"\n[5/6] Scanning for unauthorized English...")
    allowed_english = load_allowed_english(translation_dict)
    print(f"   Allowed English: {sorted(allowed_english)}")

    unauthorized = scan_for_unauthorized_english(ar_doc, allowed_english)

    if unauthorized:
        print(f"   ✗ Found {len(unauthorized)} unauthorized English words:")
        for loc, word in unauthorized[:20]:  # First 20
            print(f"      - {word} at {loc}")
    else:
        print(f"   ✓ No unauthorized English found")

    # Summary
    print(f"\n[6/6] Generating summary...")

    all_pass = (
        all(r['status'] == 'PASS' for r in structure_results) and
        alignment_result['percentage'] == 100 and
        rtl_result['percentage'] == 100 and
        len(unauthorized) == 0
    )

    print("\n" + "="*80)
    if all_pass:
        print("✅ ALL CHECKS PASSED")
        print("\nDocument is ready for delivery:")
        print(f"  - Structure matches English exactly")
        print(f"  - All cells right-aligned ({alignment_result['total']} cells)")
        print(f"  - All elements RTL-formatted ({rtl_result['total']} elements)")
        print(f"  - No unauthorized English found")
    else:
        print("⚠️  SOME CHECKS FAILED")
        print("\nIssues found:")

        if not all(r['status'] == 'PASS' for r in structure_results):
            print("  - Structure mismatch (see above)")

        if alignment_result['percentage'] != 100:
            print(f"  - {alignment_result['total'] - alignment_result['right_aligned']} cells not right-aligned")

        if rtl_result['percentage'] != 100:
            print(f"  - {rtl_result['total'] - rtl_result['rtl_formatted']} elements missing RTL formatting")

        if unauthorized:
            print(f"  - {len(unauthorized)} unauthorized English words found")

        print("\nReview details above and fix issues before delivery.")

    print("="*80)

    # Exit code
    sys.exit(0 if all_pass else 1)


if __name__ == '__main__':
    main()
