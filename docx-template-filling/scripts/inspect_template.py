#!/usr/bin/env python3
"""
Inspect DOCX template structure before modification.

This script prints complete template analysis to prevent destructive mistakes.
Always run before filling templates.

Usage:
    python scripts/inspect_template.py <template.docx>
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)


def inspect_template(doc_path):
    """
    Print complete template structure before any modifications.

    Identifies:
    - Table types and identities
    - Anchor points for content insertion
    - Headers/footers
    - Document element counts

    Prevents:
    - Modifying wrong tables
    - Missing anchor points
    - Breaking headers/footers
    - Index out-of-bounds errors
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        print(f"Error: File not found: {doc_path}")
        sys.exit(1)

    doc = Document(doc_path)

    print("=" * 70)
    print("TEMPLATE STRUCTURE ANALYSIS")
    print("=" * 70)

    # 1. High-level counts
    print(f"\nDocument Elements:")
    print(f"  Tables: {len(doc.tables)}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Sections: {len(doc.sections)}")

    # 2. Table identities
    print(f"\nTable Analysis:")
    for i, table in enumerate(doc.tables):
        first_cell = table.rows[0].cells[0].text[:60] if table.rows else ""
        print(f"  Table {i}:")
        print(f"    Size: {len(table.rows)}x{len(table.columns)}")
        print(f"    First cell: '{first_cell}...'")

        # Sample first row to identify table type
        if table.rows and table.rows[0].cells:
            first_row_text = " | ".join([c.text[:20] for c in table.rows[0].cells])
            print(f"    First row: {first_row_text}")

    # 3. Potential anchor points
    print(f"\nPotential Anchor Points:")
    anchors_found = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Common anchor patterns
        if (text.endswith(':') or
            'Answer' in text or
            'Summary' in text or
            'FILL' in text or
            'Response' in text or
            'Solution' in text or
            text in ['', '\n']):  # Empty paragraphs might be fill points

            print(f"  Para {i}: '{text}' (style: {para.style.name})")
            anchors_found += 1

            if anchors_found > 20:  # Limit output
                print(f"  ... ({len(doc.paragraphs) - i} more paragraphs)")
                break

    # 4. Headers/Footers
    print(f"\nHeaders/Footers:")
    for i, section in enumerate(doc.sections):
        header_text = section.header.paragraphs[0].text[:50] if section.header.paragraphs else "(empty)"
        footer_text = section.footer.paragraphs[0].text[:50] if section.footer.paragraphs else "(empty)"
        print(f"  Section {i}:")
        print(f"    Header: {header_text}")
        print(f"    Footer: {footer_text}")

    print("=" * 70)
    print("\nNow safe to proceed with modifications.")
    print("=" * 70)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} <template.docx>")
        sys.exit(1)

    inspect_template(sys.argv[1])
